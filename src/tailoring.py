"""LLM-powered resume and cover letter tailoring using OpenAI."""

import os
import json
import time
from pathlib import Path

from dotenv import load_dotenv
from openai import OpenAI

from .profile import load_profile, get_profile_summary
from .utils import TEMPLATES_DIR

load_dotenv()

# ── Anti-fabrication safeguard — HARDCODED, not configurable ─────────
SYSTEM_PROMPT = """You are a professional resume and cover letter writer.

CRITICAL RULES (non-negotiable):
1. You must ONLY use experience, skills, credentials, and achievements that appear in the provided base resume.
2. NEVER invent, fabricate, exaggerate, or embellish any qualification, skill, project, or achievement.
3. If the candidate lacks a required skill listed in the job description, do NOT add it. Simply omit it.
4. The base resume is the SINGLE SOURCE OF TRUTH for the candidate's background.
5. You may reorder, emphasize, and reword existing content to better match the job — but never add new claims.
6. Use professional tone, ATS-friendly formatting (standard headings, no tables, no graphics).
7. The resume MUST fit on exactly ONE page. Be concise — use short bullet points (1 line each), limit to 3-4 bullets per role, and only include the most relevant experience. Cut less relevant roles or merge them into a brief line.
8. Keep the cover letter to under 1 page.
"""


def _get_client(settings: dict) -> OpenAI:
    """Get OpenAI client."""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or api_key == "your-openai-api-key-here":
        raise ValueError("OPENAI_API_KEY not set in .env file. Please add your key.")
    timeout = settings.get("openai", {}).get("timeout", 120)
    return OpenAI(api_key=api_key, timeout=timeout)


def _get_model(settings: dict) -> str:
    return settings.get("openai", {}).get("model", "gpt-4o")


def _get_temperature(settings: dict) -> float:
    return settings.get("openai", {}).get("temperature", 0.7)


def _call_with_retry(client: OpenAI, settings: dict, prompt: str,
                     max_retries: int = 3) -> str:
    """Call the OpenAI API with retry + exponential backoff for transient failures."""
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=_get_model(settings),
                temperature=_get_temperature(settings),
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": prompt},
                ],
            )
            return response.choices[0].message.content
        except KeyboardInterrupt:
            raise
        except Exception as e:
            if attempt == max_retries - 1:
                raise
            wait = 2 ** (attempt + 1)  # 2s, 4s
            print(f"  OpenAI API error (attempt {attempt+1}/{max_retries}): {e}")
            print(f"  Retrying in {wait}s...")
            time.sleep(wait)


def load_base_resume() -> str:
    """Load and extract text from the base resume DOCX."""
    resume_path = TEMPLATES_DIR / "base_resume.docx"
    if not resume_path.exists():
        raise FileNotFoundError(
            f"Base resume not found at {resume_path}. "
            "Please place your resume as templates/base_resume.docx"
        )
    from docx import Document
    doc = Document(str(resume_path))
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def load_base_cover_letter() -> str:
    """Load and extract text from the base cover letter DOCX."""
    cl_path = TEMPLATES_DIR / "base_cover_letter.docx"
    if not cl_path.exists():
        raise FileNotFoundError(
            f"Base cover letter not found at {cl_path}. "
            "Please place your cover letter as templates/base_cover_letter.docx"
        )
    from docx import Document
    doc = Document(str(cl_path))
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def tailor_resume(job: dict, settings: dict) -> str:
    """Generate a tailored resume for a specific job.

    Returns the tailored resume as formatted text.
    """
    client = _get_client(settings)
    profile = load_profile()
    base_resume = load_base_resume()
    profile_summary = get_profile_summary(profile)

    prompt = f"""Tailor the following resume for this specific job posting.

## Job Details
- Title: {job.get('title', 'N/A')}
- Company: {job.get('company', 'N/A')}
- Location: {job.get('location', 'N/A')}

## Job Description
{job.get('description', 'No description available.')}

## Candidate Profile
{profile_summary}

## Base Resume (SOURCE OF TRUTH — only use information from here)
{base_resume}

## Instructions
1. Reorder bullet points to highlight experience most relevant to this role
2. Adjust the summary/objective to target this specific position (2-3 sentences max)
3. Emphasize skills and technologies mentioned in the job description that the candidate actually has
4. Use strong action verbs and quantify achievements where the data exists in the original
5. Format with standard headings: Summary, Experience, Education, Skills
6. CRITICAL: The resume MUST fit on exactly ONE page. To achieve this:
   - Keep the summary to 2-3 sentences
   - Limit each role to 3-4 concise bullet points (1 line each)
   - Only include the most relevant roles; omit or condense less relevant ones
   - Keep the Skills section to a single compact line or two
   - Education should be 1-2 lines total
7. DO NOT add any skills, experience, or achievements not present in the base resume
"""

    return _call_with_retry(client, settings, prompt)


def tailor_cover_letter(job: dict, settings: dict) -> str:
    """Generate a tailored cover letter for a specific job.

    Returns the tailored cover letter as formatted text.
    """
    client = _get_client(settings)
    profile = load_profile()
    base_resume = load_base_resume()
    profile_summary = get_profile_summary(profile)

    try:
        base_cl = load_base_cover_letter()
    except FileNotFoundError:
        base_cl = "(No base cover letter provided — generate from resume.)"

    prompt = f"""Write a tailored cover letter for this specific job posting.

## Job Details
- Title: {job.get('title', 'N/A')}
- Company: {job.get('company', 'N/A')}
- Location: {job.get('location', 'N/A')}

## Job Description
{job.get('description', 'No description available.')}

## Candidate Profile
{profile_summary}

## Base Resume (SOURCE OF TRUTH)
{base_resume}

## Base Cover Letter (for tone/style reference)
{base_cl}

## Instructions
1. Address the specific company and role by name
2. Connect the candidate's actual experience to the job requirements
3. Highlight 2-3 key achievements from the resume that are most relevant
4. Show enthusiasm for the company and role
5. Keep it under 1 page (roughly 300-400 words)
6. Professional but personable tone
7. DO NOT mention skills or experience not present in the base resume
8. End with a clear call to action
"""

    return _call_with_retry(client, settings, prompt)


def infer_form_answers(fields: list[dict], job: dict, settings: dict) -> dict:
    """Use LLM to infer answers for application form fields.

    Args:
        fields: List of form field dicts with keys: id, label, type, options (if select), required
        job: Job data dict
        settings: App settings

    Returns:
        Dict mapping field id to answer value
    """
    client = _get_client(settings)
    profile = load_profile()
    profile_summary = get_profile_summary(profile)

    # Include diversity info if available
    diversity = profile.get("diversity", {})
    diversity_lines = []
    for key, val in diversity.items():
        if val:
            diversity_lines.append(f"  {key}: {val}")
    diversity_text = "\n".join(diversity_lines) if diversity_lines else "  (Not provided — use 'Prefer not to answer' or 'Decline to self-identify' when available)"

    prompt = f"""You are filling out a job application form. Given the applicant's profile and the form fields below, return a JSON object mapping each field's "id" to the value that should be entered.

## Applicant Profile
{profile_summary}

## Diversity/EEO Information
{diversity_text}

## Job Being Applied For
- Title: {job.get('title', 'N/A')}
- Company: {job.get('company', 'N/A')}

## Form Fields
{json.dumps(fields, indent=2)}

## Rules
1. For text fields: provide the appropriate value from the profile
2. For select/dropdown fields: choose the EXACT option text from the available options
3. For radio/checkbox fields: choose the EXACT option label
4. For fields about salary: use the candidate's desired salary range
5. For diversity questions with no profile data: prefer "Decline to self-identify" or "Prefer not to answer" if available
6. For "How did you hear about us": use "Job Board" or similar generic option
7. Only use real information from the profile — never fabricate

Return ONLY valid JSON — no explanation, no markdown fences.
"""

    for attempt in range(3):
        try:
            response = client.chat.completions.create(
                model=_get_model(settings),
                temperature=0.3,  # Lower temperature for form filling — want precision
                messages=[
                    {"role": "system", "content": "You fill out job application forms accurately using the applicant's real profile data. Return only valid JSON."},
                    {"role": "user", "content": prompt},
                ],
            )
            break
        except KeyboardInterrupt:
            raise
        except Exception as e:
            if attempt == 2:
                raise
            wait = 2 ** (attempt + 1)
            print(f"  OpenAI API error (attempt {attempt+1}/3): {e}")
            print(f"  Retrying in {wait}s...")
            time.sleep(wait)

    text = response.choices[0].message.content.strip()
    # Strip markdown code fences if present
    if text.startswith("```"):
        text = text.split("\n", 1)[1] if "\n" in text else text[3:]
    if text.endswith("```"):
        text = text[:-3]
    text = text.strip()

    return json.loads(text)
