"""LLM-powered resume and cover letter tailoring using OpenAI."""

import os
import json
import logging
import re
import time
from pathlib import Path

from dotenv import load_dotenv
from openai import OpenAI

from ..config import load_profile, get_profile_summary
from ..utils import TEMPLATES_DIR

load_dotenv()

logger = logging.getLogger(__name__)

# -- Anti-fabrication safeguard -- HARDCODED, not configurable ---------
SYSTEM_PROMPT = """You are a professional resume and cover letter writer.

CRITICAL RULES (non-negotiable):
1. You must ONLY use experience, skills, credentials, and achievements that appear in the provided base resume.
2. NEVER invent, fabricate, exaggerate, or embellish any qualification, skill, project, or achievement.
3. If the candidate lacks a required skill listed in the job description, do NOT add it. Simply omit it.
4. The base resume is the SINGLE SOURCE OF TRUTH for the candidate's background.
5. You may reorder, emphasize, and reword existing content to better match the job — but never add new claims.
6. Use professional tone, ATS-friendly formatting (standard headings, no tables, no graphics).
7. The resume MUST fit on exactly ONE page. Be concise — use short bullet points (1 line each), limit to 3-4 bullets per role, and only include the most relevant experience. Cut less relevant roles or merge them into a brief line.
8. Keep the cover letter to under 1 page.
"""


def _get_client(settings: dict) -> OpenAI:
    """Get OpenAI client."""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or api_key == "your-openai-api-key-here":
        raise ValueError("OPENAI_API_KEY not set in .env file. Please add your key.")
    timeout = settings.get("openai", {}).get("timeout", 120)
    return OpenAI(api_key=api_key, timeout=timeout)


def _get_model(settings: dict) -> str:
    return settings.get("openai", {}).get("model", "gpt-4o")


def _get_form_model(settings: dict) -> str:
    openai_cfg = settings.get("openai", {})
    return openai_cfg.get("form_model") or openai_cfg.get("model", "gpt-4o")


def _get_temperature(settings: dict) -> float:
    return settings.get("openai", {}).get("temperature", 0.7)


def _call_with_retry(client: OpenAI, settings: dict, prompt: str,
                     max_retries: int = 3) -> str:
    """Call the OpenAI API with retry + exponential backoff for transient failures."""
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=_get_model(settings),
                temperature=_get_temperature(settings),
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": prompt},
                ],
            )
            return response.choices[0].message.content
        except KeyboardInterrupt:
            raise
        except Exception as e:
            if attempt == max_retries - 1:
                raise
            wait = 2 ** (attempt + 1)  # 2s, 4s
            logger.warning(f"OpenAI API error (attempt {attempt+1}/{max_retries}): {e}")
            logger.info(f"Retrying in {wait}s...")
            time.sleep(wait)


def load_base_resume() -> str:
    """Load and extract text from the base resume DOCX."""
    resume_path = TEMPLATES_DIR / "base_resume.docx"
    if not resume_path.exists():
        raise FileNotFoundError(
            f"Base resume not found at {resume_path}. "
            "Please place your resume as templates/base_resume.docx"
        )
    from docx import Document
    doc = Document(str(resume_path))
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def load_base_cover_letter() -> str:
    """Load and extract text from the base cover letter DOCX."""
    cl_path = TEMPLATES_DIR / "base_cover_letter.docx"
    if not cl_path.exists():
        raise FileNotFoundError(
            f"Base cover letter not found at {cl_path}. "
            "Please place your cover letter as templates/base_cover_letter.docx"
        )
    from docx import Document
    doc = Document(str(cl_path))
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def tailor_resume(job: dict, settings: dict) -> str:
    """Generate a tailored resume for a specific job.

    Returns the tailored resume as formatted text.
    """
    client = _get_client(settings)
    profile = load_profile()
    base_resume = load_base_resume()
    profile_summary = get_profile_summary(profile)

    prompt = f"""Tailor the following resume for this specific job posting.

## Job Details
- Title: {job.get('title', 'N/A')}
- Company: {job.get('company', 'N/A')}
- Location: {job.get('location', 'N/A')}

## Job Description
{job.get('description', 'No description available.')}

## Candidate Profile
{profile_summary}

## Base Resume (SOURCE OF TRUTH — only use information from here)
{base_resume}

## Instructions
1. Reorder bullet points to highlight experience most relevant to this role
2. Adjust the summary/objective to target this specific position (2-3 sentences max)
3. Emphasize skills and technologies mentioned in the job description that the candidate actually has
4. Use strong action verbs and quantify achievements where the data exists in the original
5. Format with standard headings: Summary, Experience, Education, Skills
6. CRITICAL: The resume MUST fit on exactly ONE page. To achieve this:
   - Keep the summary to 2-3 sentences
   - Limit each role to 3-4 concise bullet points (1 line each)
   - Only include the most relevant roles; omit or condense less relevant ones
   - Keep the Skills section to a single compact line or two
   - Education should be 1-2 lines total
7. DO NOT add any skills, experience, or achievements not present in the base resume
"""

    return _call_with_retry(client, settings, prompt)


def tailor_cover_letter(job: dict, settings: dict) -> str:
    """Generate a tailored cover letter for a specific job.

    Returns the tailored cover letter as formatted text.
    """
    client = _get_client(settings)
    profile = load_profile()
    base_resume = load_base_resume()
    profile_summary = get_profile_summary(profile)

    try:
        base_cl = load_base_cover_letter()
    except FileNotFoundError:
        base_cl = "(No base cover letter provided — generate from resume.)"

    prompt = f"""Write a tailored cover letter for this specific job posting.

## Job Details
- Title: {job.get('title', 'N/A')}
- Company: {job.get('company', 'N/A')}
- Location: {job.get('location', 'N/A')}

## Job Description
{job.get('description', 'No description available.')}

## Candidate Profile
{profile_summary}

## Base Resume (SOURCE OF TRUTH)
{base_resume}

## Base Cover Letter (for tone/style reference)
{base_cl}

## Instructions
1. Address the specific company and role by name
2. Connect the candidate's actual experience to the job requirements
3. Highlight 2-3 key achievements from the resume that are most relevant
4. Show enthusiasm for the company and role
5. Keep it under 1 page (roughly 300-400 words)
6. Professional but personable tone
7. DO NOT mention skills or experience not present in the base resume
8. End with a clear call to action
"""

    return _call_with_retry(client, settings, prompt)


def _match_answer_bank(label: str, saved: dict[str, str]) -> str | None:
    """Match a form field label against answer bank entries using keyword matching.

    Tries exact match first, then substring match (longer bank labels first
    so more specific patterns win, e.g. "first name" before "name").

    Returns the matched answer or None.
    """
    if not label:
        return None

    if _is_machine_generated_label(label):
        return None

    label_lower = label.lower()
    normalized_label = _normalize_label_for_match(label)

    # Exact match
    if label in saved and saved[label] != "N/A":
        return saved[label]

    # Keyword/substring match (case-insensitive)
    for q_label in sorted(saved.keys(), key=len, reverse=True):
        if saved[q_label] == "N/A":
            continue
        if q_label.lower() in label_lower:
            return saved[q_label]
        normalized_q_label = _normalize_label_for_match(q_label)
        if normalized_q_label and normalized_q_label in normalized_label:
            return saved[q_label]

    return None


def _normalize_label_for_match(text: str) -> str:
    """Loosen label matching for small OCR/ATS typos like doubled letters."""
    normalized = re.sub(r"[^a-z0-9]+", " ", (text or "").lower())
    normalized = re.sub(r"([a-z])\1+", r"\1", normalized)
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized


def _is_machine_generated_label(label: str) -> bool:
    """Return True for placeholder/internal labels like input-104."""
    import re

    normalized = (label or "").lower().strip()
    if not normalized:
        return True
    return bool(re.match(r"^(input|select|textarea|field|custom_select)[-_]?\d+$", normalized))


_STATE_NAMES = {
    "AL": "Alabama",
    "AK": "Alaska",
    "AZ": "Arizona",
    "AR": "Arkansas",
    "CA": "California",
    "CO": "Colorado",
    "CT": "Connecticut",
    "DE": "Delaware",
    "FL": "Florida",
    "GA": "Georgia",
    "HI": "Hawaii",
    "ID": "Idaho",
    "IL": "Illinois",
    "IN": "Indiana",
    "IA": "Iowa",
    "KS": "Kansas",
    "KY": "Kentucky",
    "LA": "Louisiana",
    "ME": "Maine",
    "MD": "Maryland",
    "MA": "Massachusetts",
    "MI": "Michigan",
    "MN": "Minnesota",
    "MS": "Mississippi",
    "MO": "Missouri",
    "MT": "Montana",
    "NE": "Nebraska",
    "NV": "Nevada",
    "NH": "New Hampshire",
    "NJ": "New Jersey",
    "NM": "New Mexico",
    "NY": "New York",
    "NC": "North Carolina",
    "ND": "North Dakota",
    "OH": "Ohio",
    "OK": "Oklahoma",
    "OR": "Oregon",
    "PA": "Pennsylvania",
    "RI": "Rhode Island",
    "SC": "South Carolina",
    "SD": "South Dakota",
    "TN": "Tennessee",
    "TX": "Texas",
    "UT": "Utah",
    "VT": "Vermont",
    "VA": "Virginia",
    "WA": "Washington",
    "WV": "West Virginia",
    "WI": "Wisconsin",
    "WY": "Wyoming",
    "DC": "District of Columbia",
}


def _match_option_text(value: str, options: list[str]) -> str | None:
    """Return the best visible option text for a model answer."""
    if not value or not options:
        return None

    value_lower = value.strip().lower()
    if not value_lower:
        return None

    for option in options:
        option_lower = option.lower()
        if (
            value_lower == option_lower
            or value_lower in option_lower
            or option_lower in value_lower
        ):
            return option
    return None


def _desired_work_setting_options(options: list[str], profile: dict, field_type: str) -> str | None:
    """Map a generic remote/work-setting preference to real visible option text."""
    remote_preference = (
        profile.get("preferences", {}).get("remote_preference", "") or ""
    ).strip().lower()

    desired_labels = []
    if not remote_preference or remote_preference in {"any", "either", "all", "no preference"}:
        desired_labels = ["Hybrid", "Remote", "On site"]
    else:
        if "hybrid" in remote_preference:
            desired_labels.append("Hybrid")
        if "remote" in remote_preference:
            desired_labels.append("Remote")
        if "site" in remote_preference or "office" in remote_preference or "onsite" in remote_preference:
            desired_labels.append("On site")

    matched = []
    for label in desired_labels:
        option = _match_option_text(label, options)
        if option and option not in matched:
            matched.append(option)

    if not matched:
        return None
    if field_type == "checkbox_group":
        return ", ".join(matched)
    return matched[0]


def _default_option_answer(field: dict, profile: dict) -> str | None:
    """Return a safe deterministic option when the model cannot answer."""
    options = field.get("options") or []
    if not options:
        return None

    field_type = field.get("type", "")
    label = " ".join(
        str(part or "").strip()
        for part in (field.get("label"), field.get("contextLabel"))
        if str(part or "").strip()
    ).lower()

    if "work setting" in label or "remote" in label:
        matched = _desired_work_setting_options(options, profile, field_type)
        if matched:
            return matched

    if (
        any(term in label for term in ("years of work experience", "years of experience", "experience do you have"))
        and field_type == "radio"
    ):
        return _match_option_text("None", options)

    return None


def _normalize_option_answer(field: dict, answer, profile: dict):
    """Normalize model answers to exact visible option text when possible."""
    if answer in (None, "", "N/A"):
        fallback = _default_option_answer(field, profile)
        return fallback or answer

    options = field.get("options") or []
    field_type = field.get("type", "")
    label = (field.get("label", "") or "").lower()
    answer_text = str(answer).strip()
    if not answer_text:
        return answer

    if "state" in label or "province" in label:
        expanded = _STATE_NAMES.get(answer_text.upper())
        if expanded and not options:
            return expanded
        if expanded:
            matched = _match_option_text(expanded, options)
            if matched:
                return matched

    if not options:
        return answer

    if "work setting" in label or "remote" in label:
        matched = _desired_work_setting_options(options, profile, field_type)
        if matched:
            return matched

    if field_type == "checkbox_group":
        matched_tokens = []
        for token in answer_text.replace(";", ",").split(","):
            token = token.strip()
            if not token:
                continue
            matched = _match_option_text(token, options)
            if matched and matched not in matched_tokens:
                matched_tokens.append(matched)
        if matched_tokens:
            return ", ".join(matched_tokens)

    matched = _match_option_text(answer_text, options)
    return matched or answer


def infer_form_answers(fields: list[dict], job: dict, settings: dict) -> dict:
    """Use LLM to infer answers for application form fields.

    Checks the answer bank first for previously saved answers (including
    profile-seeded entries). For fields the LLM can't answer from the profile,
    returns "N/A" and saves the question to the answer bank for the user to
    fill in later via `python -m src answers`.

    When fabricate_answers is enabled in settings, the LLM will generate answers
    for subjective questions (e.g. "What excites you about X?") based on the
    resume and cover letter content.

    Args:
        fields: List of form field dicts with keys: id, label, type, options (if select), required
        job: Job data dict
        settings: App settings

    Returns:
        Dict mapping field id to answer value
    """
    from ..db import get_connection, get_saved_answers, save_answers_batch

    client = _get_client(settings)
    profile = load_profile()
    profile_summary = get_profile_summary(profile)
    fabricate = settings.get("automation", {}).get("fabricate_answers", False)

    # Check answer bank for previously answered questions (includes profile-seeded entries)
    conn = get_connection()
    saved = get_saved_answers(conn)

    # Pre-fill from answer bank using keyword matching
    prefilled = {}
    remaining_fields = []
    for field in fields:
        fid = field["id"]
        label = field.get("label", "").strip()
        match = _match_answer_bank(label, saved)
        if match:
            prefilled[fid] = match
        else:
            remaining_fields.append(field)

    if prefilled:
        logger.info(f"Pre-filled {len(prefilled)} fields from answer bank + profile")

    # If all fields are pre-filled, return early
    if not remaining_fields:
        conn.close()
        return prefilled

    # Include diversity info if available
    diversity = profile.get("diversity", {})
    diversity_lines = []
    for key, val in diversity.items():
        if val:
            diversity_lines.append(f"  {key}: {val}")
    diversity_text = "\n".join(diversity_lines) if diversity_lines else "  (Not provided — use 'Prefer not to answer' or 'Decline to self-identify' when available)"

    # Build fabrication context if enabled
    fabrication_section = ""
    fabrication_rules = ""
    if fabricate:
        try:
            resume_text = load_base_resume()
        except FileNotFoundError:
            resume_text = ""
        try:
            cl_text = load_base_cover_letter()
        except FileNotFoundError:
            cl_text = ""

        if resume_text or cl_text:
            fabrication_section = f"""
## Resume Content (for generating answers to subjective questions)
{resume_text[:3000] if resume_text else '(not available)'}

## Cover Letter Content (for tone and motivation)
{cl_text[:2000] if cl_text else '(not available)'}
"""
            fabrication_rules = """
7. For subjective/motivational questions (e.g. "What excites you about [company]?", "Why are you interested in this role?"):
   Generate a thoughtful, professional answer based on the resume, cover letter, and job details. Keep it concise (2-3 sentences).
8. For yes/no factual questions about deadlines, accommodations, or scheduling constraints: answer "No" unless the profile says otherwise.
9. For "additional information" or open-ended optional fields: provide a brief, professional response drawing from the resume/cover letter, or leave as "N/A" if truly irrelevant.
10. STILL never fabricate skills, experience, credentials, or qualifications not in the resume.
11. If you truly cannot determine the answer, return "N/A"."""
        else:
            fabrication_rules = """
7. Only use real information from the profile — never fabricate
8. If you cannot determine the answer from the profile, return "N/A" for that field — do NOT guess or make up answers"""
    else:
        fabrication_rules = """
7. Only use real information from the profile — never fabricate
8. If you cannot determine the answer from the profile, return "N/A" for that field — do NOT guess or make up answers"""

    prompt = f"""You are filling out a job application form. Given the applicant's profile and the form fields below, return a JSON object mapping each field's "id" to the value that should be entered.

## Applicant Profile
{profile_summary}

## Diversity/EEO Information
{diversity_text}

## Job Being Applied For
- Title: {job.get('title', 'N/A')}
- Company: {job.get('company', 'N/A')}
{fabrication_section}
## Form Fields
{json.dumps([{k: v for k, v in f.items() if not k.startswith('_')} for f in remaining_fields], indent=2)}

## Rules
1. For text fields: provide the appropriate value from the profile
2. For select/dropdown fields: choose the EXACT option text from the available options
3. For radio/checkbox fields: choose the EXACT option label
4. For fields about salary: use the candidate's desired salary range
5. For diversity questions with no profile data: prefer "Decline to self-identify" or "Prefer not to answer" if available
6. For "How did you hear about us": use "Job Board" or similar generic option
{fabrication_rules}

Return ONLY valid JSON — no explanation, no markdown fences.
"""

    field_labels = [f.get("label", f.get("id", "?")) for f in fields]
    logger.info(f"Form filling: {len(fields)} fields for {job.get('company', '?')} - {job.get('title', '?')}")
    logger.debug(f"Form fields: {field_labels}")

    start_time = time.time()
    for attempt in range(3):
        try:
            logger.debug(f"OpenAI API call for form filling (attempt {attempt+1}/3)...")
            response = client.chat.completions.create(
                model=_get_form_model(settings),
                temperature=0.3,  # Lower temperature for form filling — want precision
                messages=[
                    {"role": "system", "content": "You fill out job application forms accurately using the applicant's real profile data. Return only valid JSON."},
                    {"role": "user", "content": prompt},
                ],
            )
            elapsed = time.time() - start_time
            usage = response.usage
            logger.info(
                f"Form filling API response: {elapsed:.1f}s, "
                f"{usage.prompt_tokens} prompt + {usage.completion_tokens} completion = {usage.total_tokens} tokens"
            )
            break
        except KeyboardInterrupt:
            raise
        except Exception as e:
            if attempt == 2:
                raise
            wait = 2 ** (attempt + 1)
            logger.warning(f"OpenAI API error (attempt {attempt+1}/3): {e}")
            logger.info(f"Retrying in {wait}s...")
            time.sleep(wait)

    text = response.choices[0].message.content.strip()
    # Strip markdown code fences if present
    if text.startswith("```"):
        text = text.split("\n", 1)[1] if "\n" in text else text[3:]
    if text.endswith("```"):
        text = text[:-3]
    text = text.strip()

    answers = json.loads(text)
    for field in fields:
        fid = field["id"]
        if _is_machine_generated_label(field.get("label", "")) and fid in answers:
            answers[fid] = "N/A"
        if fid in answers:
            answers[fid] = _normalize_option_answer(field, answers[fid], profile)
        if fid in prefilled:
            prefilled[fid] = _normalize_option_answer(field, prefilled[fid], profile)
    logger.debug(f"Form answers: {json.dumps(answers, indent=2)}")

    # Save N/A questions to the answer bank for the user to fill in later
    na_questions = []
    for field in remaining_fields:
        fid = field["id"]
        label = field.get("label", "").strip()
        if label and answers.get(fid) == "N/A":
            na_questions.append(label)
    if na_questions:
        save_answers_batch(conn, na_questions, source="auto")
        logger.info(f"Saved {len(na_questions)} unanswered questions to answer bank")

    conn.close()

    # Merge prefilled answers with LLM answers
    answers.update(prefilled)
    return answers
