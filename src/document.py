"""Document generation — creates DOCX and PDF versions of tailored resumes and cover letters."""

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .utils import get_application_dir


def _add_formatted_text(doc: Document, text: str):
    """Parse LLM-generated text and add it to a DOCX document with basic formatting.

    Handles markdown-style headings (##), bold (**text**), and bullet points (- or *).
    """
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Headings
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, content)
        else:
            p = doc.add_paragraph()
            _add_inline_formatting(p, stripped)


def _add_inline_formatting(paragraph, text: str):
    """Handle **bold** inline formatting within a paragraph."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def create_resume_docx(tailored_text: str, company: str, position: str) -> Path:
    """Create a DOCX resume from tailored text.

    Returns the path to the created file.
    """
    app_dir = get_application_dir(company, position)
    output_path = app_dir / "resume.docx"

    doc = Document()

    # Set default font — 10.5pt for one-page fit
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)

    # Tighten paragraph spacing for one-page resume
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(2)
    paragraph_format.line_spacing = 1.0

    # Narrow margins to maximize space
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    _add_formatted_text(doc, tailored_text)
    doc.save(str(output_path))
    return output_path


def create_cover_letter_docx(tailored_text: str, company: str, position: str) -> Path:
    """Create a DOCX cover letter from tailored text.

    Returns the path to the created file.
    """
    app_dir = get_application_dir(company, position)
    output_path = app_dir / "cover_letter.docx"

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    _add_formatted_text(doc, tailored_text)
    doc.save(str(output_path))
    return output_path


def _sanitize_for_pdf(text: str) -> str:
    """Replace Unicode characters unsupported by built-in PDF fonts with ASCII equivalents."""
    replacements = {
        "\u2013": "-",   # en-dash
        "\u2014": "--",  # em-dash
        "\u2018": "'",   # left single quote
        "\u2019": "'",   # right single quote
        "\u201c": '"',   # left double quote
        "\u201d": '"',   # right double quote
        "\u2026": "...", # ellipsis
        "\u2022": "-",   # bullet
        "\u00a0": " ",   # non-breaking space
        "\u2011": "-",   # non-breaking hyphen
        "\u00b7": "-",   # middle dot
        "\u2010": "-",   # hyphen
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    # Strip any remaining non-latin1 characters
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _pdf_add_formatted_text(pdf, text: str):
    """Parse LLM-generated text and render it to a PDF via fpdf2."""
    text = _sanitize_for_pdf(text)
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            pdf.ln(3)
            continue

        if stripped.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.cell(0, 6, stripped[4:], new_x="LMARGIN", new_y="NEXT")
        elif stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 7, stripped[3:], new_x="LMARGIN", new_y="NEXT")
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(2)
        elif stripped.startswith("# "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 8, stripped[2:], new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:]
            pdf.set_font("Helvetica", "", 10)
            pdf.cell(5, 5, "-")  # bullet (ASCII-safe for built-in fonts)
            _pdf_write_inline(pdf, content, 5)
        else:
            pdf.set_font("Helvetica", "", 10)
            _pdf_write_inline(pdf, stripped, 5)


def _pdf_write_inline(pdf, text: str, line_height: float):
    """Write text with **bold** inline formatting."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            pdf.set_font("Helvetica", "B", 10)
            pdf.write(line_height, part[2:-2])
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.write(line_height, part)
    pdf.ln(line_height)


def create_resume_pdf(tailored_text: str, company: str, position: str) -> Path:
    """Create a PDF resume directly from tailored text using fpdf2."""
    from fpdf import FPDF

    app_dir = get_application_dir(company, position)
    output_path = app_dir / "resume.pdf"

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=10)
    pdf.set_margins(13, 10, 13)  # tight margins for one-page resume

    _pdf_add_formatted_text(pdf, tailored_text)
    pdf.output(str(output_path))
    return output_path


def create_cover_letter_pdf(tailored_text: str, company: str, position: str) -> Path:
    """Create a PDF cover letter directly from tailored text using fpdf2."""
    from fpdf import FPDF

    app_dir = get_application_dir(company, position)
    output_path = app_dir / "cover_letter.pdf"

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.set_margins(25, 25, 25)  # standard margins for cover letter

    _pdf_add_formatted_text(pdf, tailored_text)
    pdf.output(str(output_path))
    return output_path


def save_application_metadata(company: str, position: str, job: dict, form_answers: dict = None):
    """Save application metadata as JSON alongside the documents."""
    app_dir = get_application_dir(company, position)
    metadata = {
        "job_title": job.get("title"),
        "company": job.get("company"),
        "location": job.get("location"),
        "job_url": job.get("url"),
        "site": job.get("site"),
        "applied_at": datetime.now().isoformat(),
        "form_answers": form_answers,
    }
    meta_path = app_dir / "application.json"
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2, ensure_ascii=False)
    return meta_path
